from django.http import JsonResponse, HttpRequest
from django.shortcuts import render
from django.db.models import Sum

from rest_framework.response import Response
from rest_framework import viewsets, permissions, status, filters
from rest_framework.decorators import action

from django_filters.rest_framework import DjangoFilterBackend, IsoDateTimeFilter, DateFilter, LookupChoiceFilter, CharFilter
from django_filters import FilterSet, DateTimeFromToRangeFilter

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from num2words import num2words
import calendar
import datetime

from apps.clients.models import *
from apps.collaterals.models import *
from .serializers import *


class CreditFilter(FilterSet):
    num_dog = CharFilter()
    client__individual_client__rnn = CharFilter()
    client__individual_client__iin = CharFilter()
    client = CharFilter()
    class Meta:
        model = Credit
        fields = ['num_dog', 'client__individual_client__rnn','client__individual_client__iin', 'status', 'client', 'credit_type']

class CreditViewSet(viewsets.ModelViewSet):
    """Регистрация, Получение, Удаление, Изменение, Частичное Изменение"""
    queryset = Credit.objects.all()
    serializer_class = CreditSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filterset_class = CreditFilter
    filter_backends = [DjangoFilterBackend,] #filters.SearchFilter]

    def get_serializer_class(self):
        if self.action == 'list' or self.action == 'retrieve':
            return CreditGetSerializer
        return self.serializer_class
    
    @staticmethod
    def fill_field(self, paragraph, field_name, field_value):
        if(field_value == None):
            field_value = "________________"
        if field_name in paragraph.text:
            run = paragraph.runs[0]  # Получаем первый Run в параграфе
            font_size = run.font.size  # Получаем размер шрифта
            paragraph.text = paragraph.text.replace(field_name, field_value)
            new_run = paragraph.runs[0]  # Получаем новый Run
            new_run.font.size = font_size  # Устанавливаем размер шрифта
            new_run.font.name = 'Times New Roman'  # Устанавливаем шрифт Times New Roman


    @action(detail=True, methods=['get'])
    def generate_credit_treatment(self, request, pk=None):
        try:
            credit = self.get_object()
        except:
            raise Response(status=status.HTTP_404_NOT_FOUND)

        # try:
        # kausar_company = Company.objects.get(reg_num=123123123)
        # kausar_company_director = IndividualClient.objects.filter(reg_number="0001").first()
        
        try:
            percent_rate = float(credit.effective_rate)
            loan_term = int(credit.period_count)
            loan_amount = float(credit.amount)
            commission_rate = 2
            days_in_first_payment = 30
            days_in_last_payment = 30
            monthly_commission_in = 0
            data = self.calculate_payment(self= self, credit=credit, percent_rate=percent_rate, loan_term=loan_term, loan_amount=loan_amount, commission_rate=commission_rate, days_in_first_payment=days_in_first_payment, days_in_last_payment=days_in_last_payment, monthly_commission_in=monthly_commission_in, with_creation=False, reset=False, credit_payment_type=1)
        except:
            return Response({"error_message":"Ошибка при расчете платежей"},status=status.HTTP_404_NOT_FOUND)        
        
        credit_client = credit.client
        
        client_docs = Docs.objects.filter(client=credit_client, identity_card_type=1).last()
        
        if(client_docs == None):
            return Response({"error_message":"У клиента нет документа удостоверяющего личность"},status=status.HTTP_404_NOT_FOUND)
        
        client_address_fact = Address.objects.filter(client=credit_client, addr_type=1).last()
        client_address_reg = Address.objects.filter(client=credit_client, addr_type=2).last()
        
        # return Response({"client_address_fact": client_address_fact, "client_address_reg": client_address_reg},status=status.HTTP_404_NOT_FOUND)
        if(client_address_fact == None or client_address_reg == None):
            return Response({"error_message":"У клиента нет адресов"},status=status.HTTP_404_NOT_FOUND)
        
        client_contact_home = Contact.objects.filter(client=credit_client, contact_type=1).last()
        client_contact_phone = Contact.objects.filter(client=credit_client, contact_type=2).last()
        
        if(client_contact_home == None or client_contact_phone == None):
            return Response({"error_message":"У клиента нет контактов"},status=status.HTTP_404_NOT_FOUND)
        
        try:
            document = Document('apps/credits/utils/template.docx')
        except:
            try:
                document = Document('kausar_nesie/apps/credits/utils/template.docx')
            except:
                return Response({"error_message":"Шаблон Документа не найден"},status=status.HTTP_404_NOT_FOUND)

        field_name_00 = "_director_full_name_"
        field_value_00 = "Директор Каусар Компани"

        field_name_01 = "_credit_protocol_number_"
        field_value_01 = credit.num_dog
        
        field_name_02 = "_credit_day_number_"
        field_value_02 = credit.date_begin.strftime("%d")
        
        date_begin_month_number = credit.date_begin.strftime("%m")
        date_begin_month_name = calendar.month_name[int(date_begin_month_number)]

        field_name_03 = "_credit_month_string_"
        field_value_03 = date_begin_month_name.capitalize()

        field_name_04 = "_credit_year_last_two_"
        field_value_04 = credit.date_begin.strftime("%y")

        field_name_05 = "_credit_year_last_one_"
        field_value_05 = credit.date_begin.strftime("%y")[1]

        field_name_06 = "_credit_amount_"
        field_value_06 = str(credit.amount).rstrip("0").rstrip(".")

        field_name_07 = "_credit_string_amount_"
        field_value_07 = num2words(str(credit.amount).rstrip("0").rstrip("."), lang='ru')

        field_name_010 = "_credit_term_"
        field_value_010 = str(credit.period_count)

        field_name_011 = "_string_credit_term_"
        field_value_011 = num2words(str(credit.period_count), lang='ru')
        

        field_name_1 = "_credit_end_day_"
        field_value_1 = credit.date_end.strftime("%d")

        date_end_month_number = credit.date_end.strftime("%m")
        date_end_month_name = calendar.month_name[int(date_end_month_number)]

        field_name_2 = "_credit_end_month_string_"
        field_value_2 = date_end_month_name.capitalize()

        field_name_3 = "_credit_end_year_last_two_"
        field_value_3 = credit.date_end.strftime("%y")

        field_name_3_1 = "_credit_payment_persent_"
        field_value_3_1 = str(credit.effective_rate).rstrip("0").rstrip(".")

        field_name_4 = "_credit_payment_persent_string_"
        field_value_4 = num2words(str(credit.effective_rate).rstrip("0").rstrip("."), lang='ru')

        field_name_5 = "_credit_payment_method_"
        field_value_5 = "Аннуитетный"

        field_name_6 = "_credit_aim_"
        field_value_6 = credit.credit_target.code

        field_name_7 = "_year_effect_rate_persent_"
        field_value_7 = str(data["gesv"])

        field_name_8 = "_year_effect_rate_persent_string_"
        field_value_8 = num2words(str(commission_rate).rstrip("0").rstrip("."), lang='ru')

        collateral = Collateral.objects.filter(client=credit.client).last()
        
        if(collateral == None):
            return Response({"error_message":"У клиента нет залога"},status=status.HTTP_404_NOT_FOUND)
        
        field_name_08 = "_collateral_type_"
        field_name_09 = "_collateral_name_"
        field_name_9 = "_collateral_date_begin_day_"
        field_name_10 = "_collateral_date_begin_month_string_"
        field_name_11 = "_collateral_date_begin_year_last_one_"
        field_name_12 = "_collateral_market_rated_company_"
        field_name_13 = "_collateral_reg_num_"
        field_name_14 = "_collateral_num_dog_"

        field_value_08 = collateral.type.name
        field_value_09 = collateral.name
        field_value_9 = collateral.date_begin.strftime("%d")
        collateral_date_begin_month_number = collateral.date_begin.strftime("%m")
        collateral_date_begin_month_name = calendar.month_name[int(collateral_date_begin_month_number)]
        field_value_10 = collateral_date_begin_month_name.capitalize()
        field_value_11 = collateral.date_begin.strftime("%y")[1]
        field_value_12 = "___________"
        field_value_13 = collateral.num_dog
        field_value_14 = collateral.num_dog

        field_name_15 = "_commission_rate_"
        field_value_15 = "2"

        field_name_16 = "_commission_string_rate_"
        field_value_16 = "Два"

        field_name_17 = "_bank_company_address_"
        field_value_17 = "г. Алматы"

        field_name_18 = "_bank_fact_address_"
        field_value_18 = "г. Алматы"

        field_name_19 = "_bank_iik_"
        field_value_19 = "123123123123123"

        field_name_20 = "_bank_kbe_"
        field_value_20 = "125121221323112"

        field_name_21 = "_country_"
        field_value_21 = "Республика Казахстан"

        field_name_22 = "_client_full_name_"
        field_value_22 = f"{credit.client.individual_client.full_name}"

        field_name_23 = "_passport_number_"
        field_value_23 = f"{client_docs.number}"

        field_name_24 = "_passport_given_place_"
        field_value_24 = f"{client_docs.issued_by}"

        field_name_25 = "_client_country_"
        field_value_25 = credit.client.individual_client.country.name

        field_name_26 = "_client_iin_"
        field_value_26 = credit.client.individual_client.iin

        field_name_27 = "_client_address_registered_"
        field_value_27 = f"{client_address_reg.areas.name}, {client_address_reg.cities.name}, {client_address_reg.street.name}, {client_address_reg.district.name}, Дом {client_address_reg.house}, Квартира {client_address_reg.flat}"

        field_name_28 = "_client_address_actual_"
        field_value_28 = f"{client_address_fact.areas.name}, {client_address_fact.cities.name}, {client_address_fact.street.name}, {client_address_fact.district.name}, Дом {client_address_fact.house}, Квартира {client_address_fact.flat}"

        field_name_29 = "_client_phone_private_"
        field_value_29 = f"{client_contact_phone.value}"

        field_name_30 = "_client_phone_home_"
        field_value_30 = f"{client_contact_home.value}"



        for paragraph in document.paragraphs:
            self.fill_field(self, paragraph, field_name_00, field_value_00)
            self.fill_field(self, paragraph, field_name_01, field_value_01)
            self.fill_field(self, paragraph, field_name_02, field_value_02)
            self.fill_field(self, paragraph, field_name_03, field_value_03)
            self.fill_field(self, paragraph, field_name_04, field_value_04)
            self.fill_field(self, paragraph, field_name_05, field_value_05)
            self.fill_field(self, paragraph, field_name_06, field_value_06)
            self.fill_field(self, paragraph, field_name_07, field_value_07)
            self.fill_field(self, paragraph, field_name_08, field_value_08)
            self.fill_field(self, paragraph, field_name_09, field_value_09)
            self.fill_field(self, paragraph, field_name_011, field_value_011)
            self.fill_field(self, paragraph, field_name_010, field_value_010)
            self.fill_field(self, paragraph, field_name_1, field_value_1)
            self.fill_field(self, paragraph, field_name_2, field_value_2)
            self.fill_field(self, paragraph, field_name_3, field_value_3)
            self.fill_field(self, paragraph, field_name_4, field_value_4)
            self.fill_field(self, paragraph, field_name_3_1, field_value_3_1)
            self.fill_field(self, paragraph, field_name_5, field_value_5)
            self.fill_field(self, paragraph, field_name_6, field_value_6)
            self.fill_field(self, paragraph, field_name_8, field_value_8)
            self.fill_field(self, paragraph, field_name_7, field_value_7)
            self.fill_field(self, paragraph, field_name_9, field_value_9)
            self.fill_field(self, paragraph, field_name_10, field_value_10)
            self.fill_field(self, paragraph, field_name_11, field_value_11)
            self.fill_field(self, paragraph, field_name_12, field_value_12)
            self.fill_field(self, paragraph, field_name_13, field_value_13)
            self.fill_field(self, paragraph, field_name_14, field_value_14)
            self.fill_field(self, paragraph, field_name_15, field_value_15)
            self.fill_field(self, paragraph, field_name_16, field_value_16)
            self.fill_field(self, paragraph, field_name_17, field_value_17)
            self.fill_field(self, paragraph, field_name_18, field_value_18)
            self.fill_field(self, paragraph, field_name_19, field_value_19)
            self.fill_field(self, paragraph, field_name_20, field_value_20)
            self.fill_field(self, paragraph, field_name_21, field_value_21)
            self.fill_field(self, paragraph, field_name_22, field_value_22)
            self.fill_field(self, paragraph, field_name_23, field_value_23)
            self.fill_field(self, paragraph, field_name_24, field_value_24)
            self.fill_field(self, paragraph, field_name_25, field_value_25)
            self.fill_field(self, paragraph, field_name_26, field_value_26)
            self.fill_field(self, paragraph, field_name_27, field_value_27)
            self.fill_field(self, paragraph, field_name_28, field_value_28)
            self.fill_field(self, paragraph, field_name_29, field_value_29)
            self.fill_field(self, paragraph, field_name_30, field_value_30)

        path = f'/credit/contracts/{field_value_01}_credit_contract.docx'

        try:
            document.save(f'kausar_nesie/media{path}')
        except:
            try:
                document.save(f'media{path}')
            except:
                return Response({"error_message":"Не удалось сохранить договор"},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        ct = CreditTreatments.objects.get_or_create(credit=credit, name=f'{field_value_01}_credit_contract.docx' , document=path)
        if(ct[0] == False):
            index = 1
        else:
            index = 0

        url = request.build_absolute_uri(ct[index].document.url)
        return Response({'status': 'ok', 'url': url}, status=status.HTTP_200_OK)

    @action(detail=False, methods=['get'])
    def test_graphic_of_payment(self, request):
        
        if(request.query_params.get('loan_term') == None):
            return Response({"error_message": "Не указан срок кредита"})
        if(request.query_params.get('loan_amount') == None):
            return Response({"error_message": "Не указана сумма кредита"})
        if(request.query_params.get('persent_rate') == None):
            return Response({"error_message": "Не указана процентная ставка"})
        if(request.query_params.get('days_in_first_payment') == None):
            temp = 30
        else:
            temp = int(request.query_params.get('days_in_first_payment'))

        if(request.query_params.get('monthly_commission_in') == None):
            mci = 0
        else:
            mci = float(request.query_params.get('monthly_commission_in'))
        
        percent_rate = float(request.query_params.get('persent_rate'))
        loan_term = int(request.query_params.get('loan_term'))
        loan_amount = float(request.query_params.get('loan_amount'))
        commission_rate = 2
        days_in_first_payment = temp
        days_in_last_payment = 30
        monthly_commission_in = mci
        if(request.query_params.get('credit_payment_type')=='1'):
            credit_payment_type = 1
        else:
            credit_payment_type = 2
        
        data = self.calculate_payment(self=self, credit=None, percent_rate=percent_rate, loan_term=loan_term, loan_amount=loan_amount, commission_rate=commission_rate, days_in_first_payment=days_in_first_payment, days_in_last_payment=days_in_last_payment, monthly_commission_in=monthly_commission_in, credit_payment_type=credit_payment_type, with_creation=False, reset=False)
        
        return render(
            request,
            'test_graphic.html',
            {
                'data':data,
            }
        )
        # return Response({"message": "График платежей успешно создан"})
        return JsonResponse(data, safe=False)

    @action(detail=True, methods=['get'])
    def generate_graphic_of_payment(self, request, pk=None):
        try:
            credit = Credit.objects.get(id=pk)
        except:
            return Response({"message": "Кредит не найден"})
        
        percent_rate = float(credit.effective_rate)
        loan_term = credit.period_count
        loan_amount = float(credit.amount)
        commission_rate = 2
        if(credit.date_sign):
            calculated_days = credit.date_begin - credit.date_sign
            days_in_first_payment = 30 + calculated_days.days
        else:
            days_in_first_payment = 30
        
        days_in_last_payment = 30
        monthly_commission_in = 0

        data = self.calculate_payment(self=self, credit=credit, percent_rate=percent_rate, loan_term=loan_term, loan_amount=loan_amount, commission_rate=commission_rate, days_in_first_payment=days_in_first_payment, days_in_last_payment=days_in_last_payment, monthly_commission_in=monthly_commission_in, with_creation=True, credit_payment_type=1)
        
        return Response({"message": "График платежей успешно создан"})
        # return JsonResponse(data, safe=False)

    @action(detail=True, methods=['get'])
    def reset_graphic_of_payment(self, request, pk=None):
        try:
            credit = Credit.objects.get(id=pk)
        except:
            return Response({"message": "Кредит не найден"})
        
        percent_rate = float(credit.effective_rate)
        loan_amount = float(credit.amount)
        commission_rate = 2
        days_in_first_payment = 30
        days_in_last_payment = 30
        monthly_commission_in = 0

        try:
            self.calculate_payment(self=self, credit=credit, percent_rate=percent_rate, loan_term=credit.period_count, loan_amount=loan_amount, commission_rate=commission_rate, days_in_first_payment=days_in_first_payment, days_in_last_payment=days_in_last_payment, monthly_commission_in=monthly_commission_in, with_creation=True, reset=True, credit_payment_type=1)
        except:
            return Response({"message": "Не удалось обнулить график платежей"})
        
        return Response({"message": "График платежей успешно обнулен"})
    
    @action(detail=True, methods=['get'])
    def regenerate_payment_schedule_by_adding_pause_months(self, request, pk=None):
        
        months = int(request.query_params.get('months'))

        if(months == None):
            return Response({"error_message": "Не указано количество месяцев"})

        if(months < 3):
            return Response({"error_message": "Количество месяцев должно быть больше 3"})
        

        credit = Credit.objects.get(id=pk)

    @action(detail=True, methods=['get'])
    def regenerate_payment_schedule_by_lowing_month_count(self, request, pk=None):
            
            if(request.query_params.get('payment_amount') == None):
                return Response({"error_message": "Не указана сумма досрочного платежа"})
            if(request.query_params.get('payment_number') == None):
                return Response({"error_message": "Не указан номер платежа"})
            if(request.query_params.get('new_month_number') == None):
                return Response({"error_message": "Не указан новый срок кредита"})

            number = int(request.query_params.get('payment_number'))
            payment_amount = float(request.query_params.get('payment_amount'))
            new_month_number = int(request.query_params.get('new_month_number'))

            credit = Credit.objects.get(id=pk)
            cps = CreditPaymentSchedule.objects.get(credit=credit, number=number)
            days_in_first_payment = 30

            # if(cps.total_payment == payment_amount):
            #     return Response({"error_message": "Сумма досрочного платежа не изменилась"})
            
            if(cps.total_payment > payment_amount):
                return Response({"error_message": "Сумма досрочного платежа больше суммы платежа"})
            
            if(number>0):
                prev_cps = CreditPaymentSchedule.objects.get(credit=credit, number=number-1)

                if number < credit.period_count:
                    new_principal_payment= round(payment_amount - (float(prev_cps.principal_remaining)*float(credit.effective_rate)/100/12),0)
                else:
                    if number == credit.period_coun:
                        new_principal_payment = round(prev_cps.principal_remaining,0)
                    else:
                        new_principal_payment = 0
                # Высчитываем Погашение возн.
                cps.amount = payment_amount
                new_commission_payment  = round((new_principal_payment*float(credit.effective_rate)/100/12)/30*days_in_first_payment,0)
                cps.commission_payment = new_commission_payment
                cps.principal_payment = new_principal_payment
                cps.total_payment = payment_amount
                cps.principal_remaining = float(prev_cps.principal_remaining) - new_principal_payment
                cps.status = PaymentStatus.objects.get(id=2)
                # print(payment_amount, new_principal_payment, new_commission_payment, float(prev_cps.principal_remaining) - new_principal_payment)
                cps.save()

                try:
                    cps_list = CreditPaymentSchedule.objects.filter(credit=credit, number__gte=number+1)
                    cps_list.delete()
                except:
                    pass

                loan_term = new_month_number + cps.number
                loan_amount = cps.principal_remaining
                principal_remaining = cps.principal_remaining
                monthly_commission_in = 0
                percent_rate = float(credit.effective_rate)
                monthly_payment = round(loan_amount * ((percent_rate*0.01/12)/(1-(1+percent_rate*0.01/12)**-(loan_term-cps.number))),0)
                
                for month in range(cps.number + 1, loan_term + 1):
                    print(month)
                    if month <= loan_term:
                        monthly_commission = round(loan_amount * monthly_commission_in, 0)
                    else:
                        monthly_commission = 0
                    # Высчитываем Погашение ОД
                    if month < loan_term:
                        principal_payment= round(monthly_payment - (principal_remaining*percent_rate/100/12),0)
                    else:
                        if month == loan_term:
                            principal_payment = round(principal_remaining,0)
                        else:
                            principal_payment = 0
                    # Высчитываем Погашение возн.
                    commission_payment  = round((principal_remaining*percent_rate/100/12)/30*days_in_first_payment,0)
                    
                    # Высчитываем Общая сумма ОД+%+ком
                    total_payment = principal_payment + commission_payment + monthly_commission
                    # Высчитываем Остаток ОД
                    principal_remaining -= principal_payment
                    
                    
                    date_to_payment = self.add_months(self, credit.date_begin, month)

                    # Создаем новый объект PaymentSchedule и заполняем его значениями из таблицы
                    # try:
                    payment_schedule = CreditPaymentSchedule.objects.create(
                        credit=credit,
                        number=month,  
                        date_to_payment=date_to_payment,  # Замените на нужное поле из таблицы, указывающее на месяц
                        date_get_payment=None,  # Замените на нужное поле из таблицы, указывающее на месяц
                        amount=0,  # Замените на нужное поле из таблицы, указывающее на сумму платежа
                        principal_payment=principal_payment,  # Замените на нужное поле из таблицы, указывающее на сумму погашения основного долга
                        commission_payment=commission_payment,  # Замените на нужное значение
                        total_payment=total_payment,  # Замените на нужное значение
                        principal_remaining=principal_remaining,  # Замените на нужное значение
                        monthly_commission=monthly_commission,  # Замените на нужное значение
                        status=PaymentStatus.objects.get(id=1),
                        penalty_commission=0,
                    )
                    payment_schedule.save()
                    # except:
                    #     pass
                        
                    

            return Response({"message": "График платежей успешно обновлен"})

    @action(detail=True, methods=['get'])
    def regenerate_payment_schedule_by_lowing_month_price(self, request, pk=None):

        if(request.query_params.get('payment_amount') == None):
            return Response({"error_message": "Не указана сумма досрочного платежа"})
        if(request.query_params.get('payment_number') == None):
            return Response({"error_message": "Не указан номер платежа"})
    
        try:
            self.update_payment_schedule(self=self, payment_amount=request.query_params.get('payment_amount'), payment_number=request.query_params.get('payment_number'), pk=pk)
        except:
            return Response({"error_message": "Не удалось обновить график платежей"},status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        return Response({"message": "График платежей успешно обновлен"})
    
    @staticmethod
    def update_payment_schedule(self, payment_amount, payment_number, pk=None):
        credit = Credit.objects.get(id=pk)
        new_first_payment_amount = payment_amount
        number = payment_number

        if(new_first_payment_amount == None):
            return Response({"error_message": "Не указана сумма досрочного платежа"})
        
        if(number == None):
            return Response({"error_message": "Не указан номер платежа"})

        cps = CreditPaymentSchedule.objects.get(credit=credit, number=number)

        if(int(number)>0):
            cps2 = CreditPaymentSchedule.objects.get(credit=credit, number=int(number)-1)
            prev_principal_remaining = float(cps2.principal_remaining)
        else:
            return Response({"error_message": "Неверный номер платежа"})
        
        if int(number) < credit.period_count:
            new_principal_payment= round(float(payment_amount) - (float(cps2.principal_remaining)*float(credit.effective_rate)/100/12),0)
        else:
            if number == credit.period_coun:
                new_principal_payment = round(cps2.principal_remaining,0)
            else:
                new_principal_payment = 0
        
        cps.amount = new_first_payment_amount
        cp  = round((float(new_principal_payment)*float(credit.effective_rate)/100/12)/30*30,0)
        cps.commission_payment = cp
        pp = float(new_first_payment_amount) - cp
        cps.principal_payment= pp
        cps.total_payment = float(new_first_payment_amount)
        cps.principal_remaining = prev_principal_remaining - pp
        cps.status = PaymentStatus.objects.get(id=2)
        cps.save()

        cpss = CreditPaymentSchedule.objects.filter(credit=credit, number__gte=int(number)+1).order_by('number')
        
        temp_principal_remaining = prev_principal_remaining - pp
        days_in_first_payment = 30
        days_in_last_payment = 30
        monthly_commission_in = 0
        percent_rate = float(credit.effective_rate)
        loan_term = credit.period_count - int(number)
        monthly_payment = round(temp_principal_remaining * ((percent_rate*0.01/12)/(1-(1+percent_rate*0.01/12)**-loan_term)),0)
        principal_remaining = temp_principal_remaining
        commission_payment = 0
        sum_income = 0
        sum_of_principal_remaining = 0
        sum_of_monthly_commission = 0
        for item in cpss:

            row = self.row_calculator(self=self, month=item.number, loan_term=loan_term, loan_amount=credit.amount, monthly_commission_in=monthly_commission_in, percent_rate=percent_rate, credit_payment_type=1, days_in_first_payment=days_in_first_payment, days_in_last_payment=days_in_last_payment, principal_remaining=principal_remaining, monthly_payment=monthly_payment)
            
            principal_payment = row['principal_payment']
            commission_payment = row['commission_payment']
            monthly_commission = row['monthly_commission']
            
            # Высчитываем Сумма %
            sum_income+=commission_payment
            # Высчитываем Остаток ОД
            principal_remaining -= principal_payment
            # Высчитываем Сумма ост. ОД
            sum_of_principal_remaining+=principal_remaining
            # Высчитываем Сумма ежемес. ком.
            sum_of_monthly_commission+=monthly_commission

            print(item.number, principal_payment, commission_payment, principal_remaining, monthly_commission)
            item.principal_payment = principal_payment
            item.commission_payment = commission_payment
            item.total_payment = principal_payment + commission_payment
            item.principal_remaining = principal_remaining
            item.monthly_commission = monthly_commission
            item.save()

        return Response({
            "number": number,
            "new_first_payment_amount": new_first_payment_amount,
            "commission_payment": cp,
            "monthly_payment": monthly_payment,
            "temp_principal_remaining": temp_principal_remaining
        })

    @staticmethod
    def add_months(self, sourcedate, months):
        month = sourcedate.month - 1 + months
        year = sourcedate.year + month // 12
        month = month % 12 + 1
        day = min(sourcedate.day, calendar.monthrange(year,month)[1])
        return datetime.date(year, month, day)
    
    @staticmethod
    def row_calculator(self, month, loan_term, loan_amount, monthly_commission_in, percent_rate, credit_payment_type, days_in_first_payment, days_in_last_payment, principal_remaining, monthly_payment):
        if month <= loan_term:
            monthly_commission = round(loan_amount * monthly_commission_in, 0)
        else:
            monthly_commission = 0
        # Высчитываем Погашение ОД
        if month < loan_term:
            if(credit_payment_type == 1):
                principal_payment= round(monthly_payment - (principal_remaining*percent_rate/100/12),0)
            else:
                principal_payment = round(loan_amount/loan_term,0)
        else:
            if month == loan_term:
                principal_payment = round(principal_remaining,0)
            else:
                principal_payment = 0
        # Высчитываем Погашение возн.
        if(month == 1):
            commission_payment  = round((principal_remaining * percent_rate/100/12)/30*days_in_first_payment,0)
        else:
            if month == loan_term:
                temp_commission = (principal_remaining * percent_rate / 100 / 12) / 30 * days_in_last_payment
            else:
                temp_commission = principal_remaining * percent_rate / 100 / 12
            commission_payment  = round(temp_commission, 0)
        
        return {
            "monthly_commission": monthly_commission,
            "principal_payment": principal_payment,
            "commission_payment": commission_payment
        }

    @staticmethod
    def calculate_payment(self, credit, percent_rate, loan_term, loan_amount, commission_rate, days_in_first_payment, days_in_last_payment, monthly_commission_in, credit_payment_type=1, with_creation=True, reset=False):
        # Расчетные данные
        monthly_commission_in/=100
        #Ежемесячный платеж
        monthly_payment = round(loan_amount * ((percent_rate*0.01/12)/(1-(1+percent_rate*0.01/12)**-loan_term)),0)
        #Сумма ком.
        commission_amount = loan_amount*commission_rate/100
        #Остаток ОД
        principal_remaining = loan_amount 
        #Изначальные данные по Сумма %,  Сумма ежемес. ком., Сумма ежемес. ком.
        sum_income = 0 #Сумма %
        sum_of_principal_remaining = principal_remaining #Сумма ост. ОД
        sum_of_monthly_commission = 0 #Сумма ежемес. ком.
        #Доход
        total_income = 0

        result = {
            "table": [],
            "gesv": 0,
            "sum_income": sum_income,
            "commission_amount": commission_amount,
            "sum_of_principal_remaining": sum_of_principal_remaining,
            "commission_rate": commission_rate,
            "total_income": total_income,
            "monthly_commission_in": monthly_commission_in,
            "sum_of_monthly_commission": sum_of_monthly_commission
        }

        result["table"].append({
            "number": 0,
            "principal_payment": 0,
            "commission_payment": 0,
            "total_payment": 0,
            "principal_remaining": principal_remaining,
            "monthly_commission": 0.0
        })
        # Создаем новый объект PaymentSchedule и заполняем его значениями из таблицы
        if(with_creation == True):
            try:
                payment_schedule = CreditPaymentSchedule.objects.get_or_create(
                    credit=credit,
                    number=0,  
                    date_to_payment=credit.date_begin,  # Замените на нужное поле из таблицы, указывающее на месяц
                    date_get_payment=None,  # Замените на нужное поле из таблицы, указывающее на месяц
                    amount=0,  # Замените на нужное поле из таблицы, указывающее на сумму платежа
                    principal_payment=0,  # Замените на нужное поле из таблицы, указывающее на сумму погашения основного долга
                    commission_payment=0,  # Замените на нужное значение
                    total_payment=0,  # Замените на нужное значение
                    principal_remaining=principal_remaining,  # Замените на нужное значение
                    monthly_commission=0.0,  # Замените на нужное значение
                    status=PaymentStatus.objects.get(id=1),
                    penalty_commission=0,
                    # status=Status.objects.get(name="Unpaid"),
                )
                payment_schedule.save()
            except:
                pass
        
        for month in range(1, int(loan_term) + 1):
            row = self.row_calculator(self, month, loan_term, loan_amount, monthly_commission_in, percent_rate, credit_payment_type, days_in_first_payment, days_in_last_payment, principal_remaining, monthly_payment)
            
            commission_payment = row["commission_payment"]
            principal_payment = row["principal_payment"]
            monthly_commission = row["monthly_commission"]

            # Высчитываем Сумма %
            sum_income+=commission_payment
            # Высчитываем Общая сумма ОД+%+ком
            total_payment = principal_payment + commission_payment + monthly_commission
            # Высчитываем Остаток ОД
            principal_remaining -= principal_payment
            # Высчитываем Сумма ост. ОД
            sum_of_principal_remaining+=principal_remaining
            # Высчитываем Сумма ежемес. ком.
            sum_of_monthly_commission+=monthly_commission
            
            if(with_creation == True):
                date_to_payment = self.add_months(self, credit.date_begin, month)

                # Создаем новый объект PaymentSchedule и заполняем его значениями из таблицы
                try:
                    payment_schedule = CreditPaymentSchedule.objects.get_or_create(
                        credit=credit,
                        number=month,  
                        date_to_payment=date_to_payment,  # Замените на нужное поле из таблицы, указывающее на месяц
                        date_get_payment=None,  # Замените на нужное поле из таблицы, указывающее на месяц
                        amount=0,  # Замените на нужное поле из таблицы, указывающее на сумму платежа
                        principal_payment=principal_payment,  # Замените на нужное поле из таблицы, указывающее на сумму погашения основного долга
                        commission_payment=commission_payment,  # Замените на нужное значение
                        total_payment=total_payment,  # Замените на нужное значение
                        principal_remaining=principal_remaining,  # Замените на нужное значение
                        monthly_commission=monthly_commission,  # Замените на нужное значение
                        status=PaymentStatus.objects.get(id=1),
                        penalty_commission=0,
                        # status=Status.objects.get(name="Unpaid"),
                    )
                    payment_schedule.save()
                except:
                    pass
            if(reset==True):
                try:
                    payment_schedule = CreditPaymentSchedule.objects.get(
                        credit=credit,
                        number=month,
                    )

                    payment_schedule.amount = 0
                    payment_schedule.principal_payment = principal_payment
                    payment_schedule.commission_payment = commission_payment
                    payment_schedule.total_payment = total_payment
                    payment_schedule.principal_remaining = principal_remaining
                    payment_schedule.monthly_commission = monthly_commission
                    payment_schedule.status = PaymentStatus.objects.get(id=1)
                    payment_schedule.penalty_commission = 0
                    payment_schedule.not_paid_amount = 0
                    payment_schedule.date_get_payment = None
                    payment_schedule.save()
                except:
                    pass 
            
            result["table"].append({
                "number": month,
                "principal_payment": principal_payment,
                "commission_payment": commission_payment,
                "total_payment": total_payment,
                "principal_remaining": principal_remaining,
                "monthly_commission": monthly_commission
            })
        total_income = sum_income + commission_amount + sum_of_monthly_commission
        result["total_income"] = total_income
        result["gesv"] = round((total_income/(sum_of_principal_remaining/loan_term))/loan_term*12*100, 1)
        result["sum_income"] = sum_income
        result["sum_of_principal_remaining"] = sum_of_principal_remaining
        return result
    
    @staticmethod
    def get_account_balance(self, client):
        try:
            account = Account.objects.get(client=client)
        except:
            return -1
        
        return account.amount

    @staticmethod
    def charge_account(self, client, amount):
        try: 
            account = Account.objects.get(client=client)
            account.amount += int(amount)
            account.save()
        except:
            return 2
        
        return 1
    
    @staticmethod
    def discharge_account(self, client, amount):
        try:
            account = Account.objects.get(client=client)
        except:
            return 2
        
        if(account.amount >= amount):
            try:
                account.amount -= int(amount)
            except:
                return 3
        else:
            return 4
        account.save()
        
        return 1

    @action(detail=True, methods=['get'])
    def recalculate_payment_day(self, request, pk=None):
        credit = self.get_object()
        payment_number = request.query_params.get('payment_number')
        payment_date = request.query_params.get('payment_date')
        
        if(payment_date == None):
            payment_date = datetime.date.today()
        else:
            payment_date = datetime.datetime.strptime(payment_date, "%Y-%m-%d").date()
        
        if(payment_number == None):
            return Response({"error": "payment_number is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            credit_payment_schedule = CreditPaymentSchedule.objects.get(credit=credit, number=payment_number)
        except:
            return Response({"error": "Не получилось найти график оплаты кредита"}, status=status.HTTP_400_BAD_REQUEST)
        
        if(credit_payment_schedule.status.id == 2):
            return Response({"error": "Уже оплачено клиентом"}, status=status.HTTP_400_BAD_REQUEST)
        
        days = (payment_date - credit_payment_schedule.date_to_payment).days

        penalty_amount = (credit_payment_schedule.total_payment - credit_payment_schedule.amount)
        
        if(days > 0 and days <= 90):
            credit_payment_schedule.penalty_commission = float(penalty_amount  * days) * 0.5/100
        if(days > 90):
            credit_payment_schedule.penalty_commission =  (float(penalty_amount  * 90) * 0.5/100) + (float(penalty_amount  * (days - 90)) * 0.03/100)

        if(credit.is_affiliated == False):
            credit.client.category_type = CategoryType.objects.get(code="1")

            if(days >= 1 and days < 30):
                print("2")
                credit.client.category_type = CategoryType.objects.get(code="2")
            elif(days >= 30 and days < 60):
                print("3")
                credit.client.category_type = CategoryType.objects.get(code="3")
            elif(days >= 60 and days < 90):
                print("4")
                credit.client.category_type = CategoryType.objects.get(code="4")
            elif(days >= 90 and days < 180):
                print("5")
                credit.client.category_type = CategoryType.objects.get(code="5")
            elif(days >= 180):
                print("6")
                credit.client.category_type = CategoryType.objects.get(code="6")

        credit.client.save()
        credit.save()
        
        credit_payment_schedule.save()
        
        return Response({"success": "Успешно пересчитана сумма оплата", "days": days}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'])
    def payment(self, request, pk=None):
        credit = self.get_object()
        payment_number = request.query_params.get('payment_number')
        payment_date = request.query_params.get('payment_date')

        if(payment_date == None):
            payment_date = datetime.date.today()
        else:
            payment_date = datetime.datetime.strptime(payment_date, "%Y-%m-%d")
        try:
            account = Account.objects.get(client=credit.client)
        except:
            return Response({"error": "Не получилось найти счёт клиента"}, status=status.HTTP_400_BAD_REQUEST)  
          
        if(payment_number == None):
            return Response({"error": "payment_number is required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            credit_payment_schedule = CreditPaymentSchedule.objects.get(credit=credit, number=payment_number)
        except:
            return Response({"error": "Не получилось найти график оплаты кредита"}, status=status.HTTP_400_BAD_REQUEST)

        if(credit_payment_schedule.status.id == 2):
            return Response({"error": "Уже оплачено клиентом"}, status=status.HTTP_400_BAD_REQUEST)
        
        balance = self.get_account_balance(self, credit.client)

        if(balance == -1):
            return Response({"error": "Не получилось найти счёт клиента"}, status=status.HTTP_400_BAD_REQUEST)
        
        if(balance == 0):
            return Response({"error": "На счёте клиента нет средств"}, status=status.HTTP_400_BAD_REQUEST)

        total_payment = credit_payment_schedule.total_payment + credit_payment_schedule.penalty_commission
        if(balance < total_payment):
            account_action_status = self.discharge_account(self, credit.client, balance)

            if(account_action_status == 2):
                return Response({"error": "Не получилось найти счёт клиента"}, status=status.HTTP_400_BAD_REQUEST)

            credit_payment_schedule.not_paid_amount = total_payment - int(balance)
            credit_payment_schedule.amount = balance
            credit_payment_schedule.date_get_payment = payment_date
            
            if(credit_payment_schedule.date_to_payment < payment_date):

                next_credit_payment_schedule = CreditPaymentSchedule.objects.get(credit=credit, number=int(payment_number)+1)
                next_credit_payment_schedule.penalty_commission = next_credit_payment_schedule.total_payment - int(balance) + credit_payment_schedule.penalty_commission
                next_credit_payment_schedule.total_payment += next_credit_payment_schedule.total_payment - int(balance)
                credit_payment_schedule.status = PaymentStatus.objects.get(id=3)
                
                next_credit_payment_schedule.save()
            else:
                credit_payment_schedule.status = PaymentStatus.objects.get(id=6)

            credit_payment_schedule.save()
        else:
            account_action_status = self.discharge_account(self, credit.client, total_payment)
            
            if(account_action_status == 2):
                return Response({"error": "Не получилось найти счёт клиента"}, status=status.HTTP_400_BAD_REQUEST)
            
            credit_payment_schedule.status = PaymentStatus.objects.get(id=2)
            credit_payment_schedule.date_get_payment = datetime.datetime.now()
            credit_payment_schedule.amount = credit_payment_schedule.total_payment
            credit_payment_schedule.save()
        
        return Response({"success": "Оплата произошла"}, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['get'])
    def payment_of_penalty(self, request, pk=None):


    @action(detail=False, methods=['get'])
    def get_last_num_reg(self, request):
        last_object = Credit.objects.last()
        if last_object:
            value = last_object.id
        else: value = 0

        return Response({'last_num_reg': value})

class CreditTreatmentViewSet(viewsets.ModelViewSet):
    queryset = CreditTreatments.objects.all()
    serializer_class = CreditTreatmentsSerializer
    permission_classes = [permissions.AllowAny]

class CreditPaymentScheduleFilter(FilterSet):
    class Meta:
        model = CreditPaymentSchedule
        fields = ['credit',]

class CreditPaymentScheduleViewSet(viewsets.ModelViewSet):
    queryset = CreditPaymentSchedule.objects.all()
    serializer_class = CreditPaymentScheduleSerializer
    permission_classes = [permissions.AllowAny]
    filterset_class = CreditPaymentScheduleFilter
    filter_backends = [DjangoFilterBackend,]
    pagination_class = None

    @action(detail=False, methods=['get'])
    def table(self, request):
        credit_id = request.query_params.get('credit_id')
        queryset = self.queryset.filter(credit=credit_id).order_by('id')
        serializer_class = self.serializer_class(queryset, many=True)
        
        return render(
            request,
            'table.html',
            {
                'data':serializer_class.data,
            }
        )
    
    

class CreditLineViewSet(viewsets.ModelViewSet):
    queryset = CreditLine.objects.all()
    serializer_class = CreditLineSerializer
    permission_classes = [permissions.AllowAny]

    @action(detail=False, methods=['get'])
    def get_last_num_reg(self, request):
        last_object = CreditLine.objects.last()
        if last_object:
            value = last_object.id
        else: value = 0

        return Response({'last_num_reg': value})
    
    @action(detail=True, methods=['post'])
    def add_credit(self, request, pk=None):
        credit_line = CreditLine.objects.get(id=pk)
        credits = credit_line.credits.all()

        try:
            credit = Credit.objects.get(id=request.data['credit_id'])
        except:
            return Response({"error_message":"Не удалось найти кредит"},status=status.HTTP_400_BAD_REQUEST)
        
        if credit.client != credit_line.client:
            return Response({"error_message":"Клиент не совпадает с клиентом кредитной линии"},status=status.HTTP_400_BAD_REQUEST)

        if credit_line.amount < credit.amount:
            return Response({"error_message":"Сумма кредит превосходить допустимую сумму кредитной линии "},status=status.HTTP_400_BAD_REQUEST)
        
        if credit.is_line == False:
            return Response({"error_message":"Кредит не имеет признаки кредитной линией"},status=status.HTTP_400_BAD_REQUEST)
        
        if credit in credits:
            return Response({"error_message":"Кредит уже находитсяв этой кредитной линии"},status=status.HTTP_400_BAD_REQUEST)
        
        if credits.count() > 0:
            total_amount = credits.aggregate(total_amount=Sum('amount'))['total_amount']
            if total_amount + credit.amount > credit_line.amount:
                return Response({"error_message":"Сумма кредит превосходить допустимую сумму кредитной линии "},status=status.HTTP_400_BAD_REQUEST)

        credit_line.credits.add(credit)
        
        return Response(status=status.HTTP_202_ACCEPTED)


class RequisiteViewSet(viewsets.ModelViewSet):
    queryset = Requisite.objects.all()
    serializer_class = RequisiteSerializer
    permission_classes = [permissions.AllowAny]